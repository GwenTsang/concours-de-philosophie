import json
import re
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Union

# ---------- Normalisation & parsing markdown inline ----------

def normalize_text(s: Optional[str]) -> str:
    if not s:
        return ""
    # Normalise les fins de ligne
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    # Corrige les échappements fréquents du JSON fourni
    s = s.replace("\\-", "-")
    s = s.replace("\\<\\<", "«").replace("\\>\\>", "»")
    # Double backslashes résiduels
    s = s.replace("\\\"", "\"")
    # Réduit espaces superflus autour de sauts de ligne (ne touche pas au début de ligne)
    s = re.sub(r"[ \t]+\n", "\n", s)
    # Ne pas strip() complètement pour préserver les tabs de tête
    return s.strip("\n")

Token = Tuple[str, Optional[str]]  # (texte, style) style in {None, 'b','i','s'}

_inline_re = re.compile(
    r"(~~.+?~~|\*\*.+?\*\*|\*.+?\*)",
    flags=re.DOTALL
)

def parse_inline_md(text: str) -> List[Token]:
    """
    Convertit **gras**, *italique*, ~~barré~~ en tokens (pour DOCX/ODT).
    NB: volontairement simple (pas d’imbrication complexe).
    """
    tokens: List[Token] = []
    pos = 0
    for m in _inline_re.finditer(text):
        if m.start() > pos:
            tokens.append((text[pos:m.start()], None))
        frag = m.group(0)
        if frag.startswith("**") and frag.endswith("**"):
            tokens.append((frag[2:-2], "b"))
        elif frag.startswith("*") and frag.endswith("*"):
            tokens.append((frag[1:-1], "i"))
        elif frag.startswith("~~") and frag.endswith("~~"):
            tokens.append((frag[2:-2], "s"))
        else:
            tokens.append((frag, None))
        pos = m.end()
    if pos < len(text):
        tokens.append((text[pos:], None))
    return tokens

def md_to_plain_paragraphs(s: str) -> List[str]:
    """
    On conserve les paragraphes (lignes vides = séparateurs) et
    SURTOUT les tabulations de début de paragraphe.
    """
    s = normalize_text(s)
    if not s:
        return []
    # Découpe par doubles sauts
    blocks = re.split(r"\n\s*\n", s)

    clean_blocks: List[str] = []
    for b in blocks:
        if not b:
            continue
        # Supprime uniquement les \n superflus en bordure,
        # conserve le leading whitespace (dont \t) à l'intérieur.
        b = re.sub(r"^\n+", "", b)
        b = re.sub(r"\n+$", "", b)
        # Supprime les espaces de fin de ligne, pas le début
        lines = [ln.rstrip() for ln in b.split("\n")]
        block = "\n".join(lines)
        if block.strip():  # garde les blocs non vides
            clean_blocks.append(block)
    return clean_blocks

# ---------- Extraction ciblée pour ton schéma ----------

def extract_structure(data: Dict) -> Dict:
    """
    Récupère uniquement :
    - en-tête : niveau+année, note, sujet
    - corps par blocs : introduction, I, transition_1, II, transition_2, III, conclusion
    Ignore les champs manquants.

    NB:
    - 'annonce_du_plan' n'est plus une section séparée ; si présent, son contenu
      est fusionné à la fin de l'introduction, AVEC saut de ligne + Tab d'alinéa.
    """
    TAB_INDENT = "\t"  # Tabulation explicite pour l’alinéa demandé

    header = {
        "concours": " ".join([v for v in [data.get("niveau"), data.get("annee")] if v]),
        "note": data.get("note"),
        "sujet": data.get("sujet") or data.get("title") or data.get("titre"),
    }

    # Pré-traiter introduction + annonce du plan
    intro_raw = data.get("introduction")
    annonce_raw = data.get("annonce_du_plan")

    intro_paragraphs: List[str] = []
    if isinstance(intro_raw, str) and intro_raw.strip():
        intro_paragraphs = md_to_plain_paragraphs(intro_raw)

    if isinstance(annonce_raw, str) and annonce_raw.strip():
        plan_txt = normalize_text(annonce_raw).strip("\n")
        if intro_paragraphs:
            # Fusion : SAUT DE LIGNE + Tabulation avant l’annonce
            intro_paragraphs[-1] = intro_paragraphs[-1].rstrip() + f"\n{TAB_INDENT}{plan_txt}"
        else:
            # S’il n’y a pas d’introduction, on crée un paragraphe unique avec Tab initiale
            intro_paragraphs = [f"{TAB_INDENT}{plan_txt}"]

    # Remplacement des intitulés P1/P2/P3 -> I/II/III
    blocks_order = [
        ("Introduction", intro_paragraphs),               # déjà traité ci-dessus
        ("I", data.get("partie_1")),
        ("Transition 1", data.get("transition_1")),
        ("II", data.get("partie_2")),
        ("Transition 2", data.get("transition_2")),
        ("III", data.get("partie_3")),
        ("Conclusion", data.get("conclusion")),
    ]

    blocks: List[Tuple[str, List[str]]] = []
    for label, raw in blocks_order:
        if isinstance(raw, list):
            paragraphs = [p for p in raw if isinstance(p, str) and p.strip()]
        else:
            paragraphs = md_to_plain_paragraphs(raw) if (isinstance(raw, str) and raw.strip()) else []
        if paragraphs:
            blocks.append((label, paragraphs))

    return {"header": header, "blocks": blocks}

# ---------- Export Markdown ----------

def export_markdown(struct: Dict, md_path: Path, rawtext: bool = False) -> None:
    H = struct["header"]
    lines: List[str] = []
    # En-tête centré (HTML pour le centrage)
    lines.append('<div align="center">')
    if H.get("sujet"):
        lines.append(f"<h1>{H['sujet']}</h1>")
    if H.get("concours"):
        lines.append(f"<h2>{H['concours']}</h2>")
    if H.get("note"):
        lines.append(f"<h3>Note : {H['note']}</h3>")
    lines.append("</div>")
    lines.append("")

    # Corps
    for label, paragraphs in struct["blocks"]:
        is_transition = bool(re.match(r"Transition\s*\d+", label, flags=re.IGNORECASE))

        # Balises centrées sauf en mode rawtext (où on ne met AUCUNE balise)
        if not is_transition and not rawtext:
            lines.append(f'<div align="center"><h2>{label}</h2></div>')
            lines.append("")

        # Paragraphes : écrits tels quels pour préserver les Tabs de tête
        for p in paragraphs:
            lines.append(p)
            lines.append("")  # saut de ligne
        lines.append("")  # séparation douce entre blocs

    md_path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")

# ---------- Export DOCX ----------

def _docx_add_line_with_leading_tabs(paragraph, line: str, style_key: Optional[str]):
    """
    Ajoute une ligne dans un paragraphe DOCX en respectant les tabulations initiales,
    puis le texte (avec style éventuel).
    """
    m = re.match(r"^\t+", line)
    tabs = len(m.group(0)) if m else 0
    content = line[tabs:]

    # Ajoute les tabs comme runs séparés (sans style)
    if tabs:
        for _ in range(tabs):
            paragraph.add_run("\t")

    # Ajoute le contenu stylé
    run = paragraph.add_run(content)
    if style_key == "i":
        run.italic = True
    elif style_key == "b":
        run.bold = True
    elif style_key == "s":
        run.font.strike = True

def export_docx(struct: Dict, docx_path: Path, rawtext: bool = False) -> None:
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ModuleNotFoundError:
        raise ImportError(
            "The 'python-docx' package is required. Install it with:\n"
            "    pip install python-docx"
        )

    FIRST_LINE_CM = 0.75  # retrait de première ligne (~0,75 cm)

    doc = Document()
    H = struct["header"]

    # Sujet (gros, centré)
    if H.get("sujet"):
        p = doc.add_paragraph()
        run = p.add_run(H["sujet"])
        run.bold = True
        run.font.size = Pt(22)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Concours (niveau + année)
    if H.get("concours"):
        p = doc.add_paragraph()
        r = p.add_run(H["concours"])
        r.bold = True
        r.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Note
    if H.get("note"):
        p = doc.add_paragraph()
        r = p.add_run(f"Note : {H['note']}")
        r.bold = True
        r.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ligne vide après l’en-tête
    doc.add_paragraph("")

    # Corps
    for label, paragraphs in struct["blocks"]:
        is_transition = bool(re.match(r"Transition\s*\d+", label, flags=re.IGNORECASE))

        # Titre de section centré (sans retrait) sauf pour les transitions et en mode rawtext
        if not is_transition and not rawtext:
            h = doc.add_paragraph()
            hr = h.add_run(label)
            hr.bold = True
            hr.font.size = Pt(14)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for para in paragraphs:
            para = normalize_text(para)
            p = doc.add_paragraph()
            # Justification du corps du texte
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(FIRST_LINE_CM)

            # Respecte les sauts internes et les tabs de tête
            for i, line in enumerate(para.split("\n")):
                if i > 0:
                    p.add_run("\n")
                tokens = parse_inline_md(line)
                if not tokens:
                    _docx_add_line_with_leading_tabs(p, line, None)
                else:
                    for j, (text_frag, style) in enumerate(tokens):
                        if j == 0:
                            _docx_add_line_with_leading_tabs(p, text_frag, style)
                        else:
                            r = p.add_run(text_frag)
                            if style == "i":
                                r.italic = True
                            elif style == "b":
                                r.bold = True
                            elif style == "s":
                                r.font.strike = True

        # Espace entre blocs
        doc.add_paragraph("")

    doc.save(str(docx_path))

# ---------- Export ODT ----------

def export_odt(struct: Dict, odt_path: Path, rawtext: bool = False) -> None:
    from odf.opendocument import OpenDocumentText
    from odf import text, style

    FIRST_LINE_CM = "0.75cm"  # retrait première ligne

    doc = OpenDocumentText()

    # Styles
    h1 = style.Style(name="H1", family="paragraph")
    h1.addElement(style.ParagraphProperties(attributes={"textalign": "center"}))
    h1.addElement(style.TextProperties(attributes={"fontsize": "22pt", "fontweight": "bold"}))
    doc.styles.addElement(h1)

    h2 = style.Style(name="H2", family="paragraph")
    h2.addElement(style.ParagraphProperties(attributes={"textalign": "center"}))
    h2.addElement(style.TextProperties(attributes={"fontsize": "16pt", "fontweight": "bold"}))
    doc.styles.addElement(h2)

    h3 = style.Style(name="H3", family="paragraph")
    h3.addElement(style.ParagraphProperties(attributes={"textalign": "center"}))
    h3.addElement(style.TextProperties(attributes={"fontsize": "14pt", "fontweight": "bold"}))
    doc.styles.addElement(h3)

    hsec = style.Style(name="HSEC", family="paragraph")
    # Balises centrées
    hsec.addElement(style.ParagraphProperties(attributes={"textalign": "center"}))
    hsec.addElement(style.TextProperties(attributes={"fontsize": "14pt", "fontweight": "bold"}))
    doc.styles.addElement(hsec)

    # Corps des paragraphes : retrait + JUSTIFY
    pstyle = style.Style(name="P", family="paragraph")
    pstyle.addElement(style.ParagraphProperties(attributes={"textindent": FIRST_LINE_CM, "textalign": "justify"}))
    pstyle.addElement(style.TextProperties(attributes={"fontsize": "11pt"}))
    doc.styles.addElement(pstyle)

    ist = style.Style(name="I", family="text")
    ist.addElement(style.TextProperties(attributes={"fontstyle": "italic"}))
    doc.styles.addElement(ist)

    bst = style.Style(name="B", family="text")
    bst.addElement(style.TextProperties(attributes={"fontweight": "bold"}))
    doc.styles.addElement(bst)

    sst = style.Style(name="S", family="text")
    sst.addElement(style.TextProperties(attributes={"textlinethroughstyle": "solid"}))
    doc.styles.addElement(sst)

    H = struct["header"]

    # Sujet / Concours / Note (non indentés)
    if H.get("sujet"):
        doc.text.addElement(text.P(stylename=h1, text=H["sujet"]))
    if H.get("concours"):
        doc.text.addElement(text.P(stylename=h2, text=H["concours"]))
    if H.get("note"):
        doc.text.addElement(text.P(stylename=h3, text=f"Note : {H['note']}"))

    # Ligne vide
    doc.text.addElement(text.P(text=""))

    # Sections
    for label, paragraphs in struct["blocks"]:
        is_transition = bool(re.match(r"Transition\s*\d+", label, flags=re.IGNORECASE))
        if not is_transition and not rawtext:
            doc.text.addElement(text.P(stylename=hsec, text=label))
        else:
            # pas de titre pour la transition OU en mode rawtext
            doc.text.addElement(text.P(text=""))

        for para in paragraphs:
            para = normalize_text(para)
            # Chaque ligne du paragraphe
            for i, line in enumerate(para.split("\n")):
                p = text.P(stylename=pstyle)
                # Gère les tabs de tête
                m = re.match(r"^\t+", line)
                tabs = len(m.group(0)) if m else 0
                rest = line[tabs:]

                for _ in range(tabs):
                    p.addElement(text.Tab())

                # Traite le markdown inline pour le reste
                for frag, style_key in parse_inline_md(rest):
                    if style_key == "i":
                        span = text.Span(stylename=ist, text=frag)
                    elif style_key == "b":
                        span = text.Span(stylename=bst, text=frag)
                    elif style_key == "s":
                        span = text.Span(stylename=sst, text=frag)
                    else:
                        span = text.Span(text=frag)
                    p.addElement(span)

                doc.text.addElement(p)

        # séparation
        doc.text.addElement(text.P(text=""))

    doc.save(str(odt_path))

# ---------- Point d’entrée ----------

def export_from_json(
    json_path: Union[str, Path],
    out_dir: Union[str, Path] = ".",
    basename: Optional[str] = None,
    rawtext: bool = False
):
    json_path = Path(json_path)
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    data = json.loads(json_path.read_text(encoding="utf-8"))
    struct = extract_structure(data)

    base = basename or (data.get("sujet") or json_path.stem)
    safe = re.sub(r"[^\w\-]+", "_", base).strip("_") or "dissertation"

    md_path = out_dir / f"{safe}.md"
    docx_path = out_dir / f"{safe}.docx"
    odt_path = out_dir / f"{safe}.odt"

    export_markdown(struct, md_path, rawtext=rawtext)
    export_docx(struct, docx_path, rawtext=rawtext)
    export_odt(struct, odt_path, rawtext=rawtext)

    return md_path, docx_path, odt_path


if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser(description="Exporter (MD/DOCX/ODT) une dissertation CAPES depuis un JSON.")
    ap.add_argument("json_path", help="Chemin du fichier JSON")
    ap.add_argument("-o", "--out-dir", default=".", help="Dossier de sortie")
    ap.add_argument("-b", "--basename", default=None, help="Nom de base des fichiers de sortie")
    # Option rawtext : supporte -rawtext (forme courte non standard) et --rawtext
    ap.add_argument("-rawtext", "--rawtext", action="store_true",
                    help="Exporter sans balises de sections (Introduction, I, II, III, Conclusion).")
    args = ap.parse_args()

    md, docx, odt = export_from_json(args.json_path, args.out_dir, args.basename, rawtext=args.rawtext)
    print("Fichiers générés :")
    print(" -", md)
    print(" -", docx)
    print(" -", odt)
