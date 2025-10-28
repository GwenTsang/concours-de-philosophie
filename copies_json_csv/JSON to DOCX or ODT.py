import json
import re
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Union

# ---------- Normalisation & parsing markdown inline ----------

def normalize_text(s: Optional[str]) -> str:
    if not s:
        return ""
    # Normalise les fins de ligne
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    # Corrige les échappements fréquents du JSON fourni
    s = s.replace("\\-", "-")
    s = s.replace("\\<\\<", "«").replace("\\>\\>", "»")
    # Double backslashes résiduels
    s = s.replace("\\\"", "\"")
    # Réduit espaces superflus autour de sauts de ligne
    s = re.sub(r"[ \t]+\n", "\n", s)
    return s.strip()

Token = Tuple[str, Optional[str]]  # (texte, style) style in {None, 'b','i','s'}

_inline_re = re.compile(
    r"(~~.+?~~|\*\*.+?\*\*|\*.+?\*)",
    flags=re.DOTALL
)

def parse_inline_md(text: str) -> List[Token]:
    """
    Convertit **gras**, *italique*, ~~barré~~ en tokens (pour DOCX/ODT).
    NB: volontairement simple (pas d’imbrication complexe).
    """
    tokens: List[Token] = []
    pos = 0
    for m in _inline_re.finditer(text):
        if m.start() > pos:
            tokens.append((text[pos:m.start()], None))
        frag = m.group(0)
        if frag.startswith("**") and frag.endswith("**"):
            tokens.append((frag[2:-2], "b"))
        elif frag.startswith("*") and frag.endswith("*"):
            tokens.append((frag[1:-1], "i"))
        elif frag.startswith("~~") and frag.endswith("~~"):
            tokens.append((frag[2:-2], "s"))
        else:
            tokens.append((frag, None))
        pos = m.end()
    if pos < len(text):
        tokens.append((text[pos:], None))
    return tokens

def md_to_plain_paragraphs(s: str) -> List[str]:
    """
    On conserve les paragraphes (lignes vides = séparateurs).
    """
    s = normalize_text(s)
    # Découpe par doubles sauts
    blocks = re.split(r"\n\s*\n", s) if s else []
    # Nettoie chaque bloc
    return [b.strip() for b in blocks if b.strip()]

# ---------- Extraction ciblée pour ton schéma ----------

def extract_structure(data: Dict) -> Dict:
    """
    Récupère uniquement :
    - en-tête : niveau+année, note, sujet
    - corps par blocs : introduction, annonce_du_plan, partie_1, transition_1, partie_2, transition_2, partie_3, conclusion
    Ignore les champs manquants.
    """
    header = {
        "concours": " ".join([v for v in [data.get("niveau"), data.get("annee")] if v]),
        "note": data.get("note"),
        "sujet": data.get("sujet") or data.get("title") or data.get("titre"),
    }

    blocks_order = [
        ("Introduction", "introduction"),
        ("Annonce du plan", "annonce_du_plan"),
        ("P1", "partie_1"),
        ("Transition 1", "transition_1"),
        ("P2", "partie_2"),
        ("Transition 2", "transition_2"),
        ("P3", "partie_3"),
        ("Conclusion", "conclusion"),
    ]

    blocks: List[Tuple[str, List[str]]] = []
    for label, key in blocks_order:
        raw = data.get(key)
        if raw and isinstance(raw, str) and raw.strip():
            paragraphs = md_to_plain_paragraphs(raw)
            if paragraphs:
                blocks.append((label, paragraphs))

    return {"header": header, "blocks": blocks}

# ---------- Export Markdown ----------

def export_markdown(struct: Dict, md_path: Path) -> None:
    H = struct["header"]
    lines: List[str] = []
    # En-tête centré (HTML pour le centrage)
    lines.append('<div align="center">')
    if H.get("sujet"):
        lines.append(f"<h1>{H['sujet']}</h1>")
    if H.get("concours"):
        lines.append(f"<h2>{H['concours']}</h2>")
    if H.get("note"):
        lines.append(f"<h3>Note : {H['note']}</h3>")
    lines.append("</div>")
    lines.append("")

    # Corps : chaque bloc avec un interligne
    for label, paragraphs in struct["blocks"]:
        lines.append(f"## {label}")
        lines.append("")
        # On laisse le markdown d’origine tel quel (les * resteront italiques)
        for p in paragraphs:
            lines.append(p)
            lines.append("")  # saut de ligne
        lines.append("")  # séparation de blocs

    md_path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")

# ---------- Export DOCX ----------

def export_docx(struct: Dict, docx_path: Path) -> None:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    H = struct["header"]

    # Sujet (gros, centré)
    if H.get("sujet"):
        p = doc.add_paragraph()
        run = p.add_run(H["sujet"])
        run.bold = True
        run.font.size = Pt(22)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Concours (niveau + année)
    if H.get("concours"):
        p = doc.add_paragraph()
        r = p.add_run(H["concours"])
        r.bold = True
        r.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Note
    if H.get("note"):
        p = doc.add_paragraph()
        r = p.add_run(f"Note : {H['note']}")
        r.bold = True
        r.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ligne vide après l’en-tête
    doc.add_paragraph("")

    # Corps
    for label, paragraphs in struct["blocks"]:
        # Titre de section
        h = doc.add_paragraph()
        hr = h.add_run(label)
        hr.bold = True
        hr.font.size = Pt(14)

        for para in paragraphs:
            para = normalize_text(para)
            p = doc.add_paragraph()
            # on segmente par lignes pour garder les sauts internes éventuels
            for i, line in enumerate(para.split("\n")):
                if i > 0:
                    p.add_run("\n")
                for text, style in parse_inline_md(line):
                    r = p.add_run(text)
                    if style == "i":
                        r.italic = True
                    elif style == "b":
                        r.bold = True
                    elif style == "s":
                        r.font.strike = True

        # Saut de ligne entre blocs
        doc.add_paragraph("")

    doc.save(str(docx_path))

# ---------- Export ODT ----------

def export_odt(struct: Dict, odt_path: Path) -> None:
    from odf.opendocument import OpenDocumentText
    from odf import text, style

    doc = OpenDocumentText()

    # Styles
    h1 = style.Style(name="H1", family="paragraph")
    h1.addElement(style.ParagraphProperties(attributes={"textalign": "center"}))
    h1.addElement(style.TextProperties(attributes={"fontsize": "22pt", "fontweight": "bold"}))
    doc.styles.addElement(h1)

    h2 = style.Style(name="H2", family="paragraph")
    h2.addElement(style.ParagraphProperties(attributes={"textalign": "center"}))
    h2.addElement(style.TextProperties(attributes={"fontsize": "16pt", "fontweight": "bold"}))
    doc.styles.addElement(h2)

    h3 = style.Style(name="H3", family="paragraph")
    h3.addElement(style.ParagraphProperties(attributes={"textalign": "center"}))
    h3.addElement(style.TextProperties(attributes={"fontsize": "14pt", "fontweight": "bold"}))
    doc.styles.addElement(h3)

    hsec = style.Style(name="HSEC", family="paragraph")
    hsec.addElement(style.TextProperties(attributes={"fontsize": "14pt", "fontweight": "bold"}))
    doc.styles.addElement(hsec)

    pstyle = style.Style(name="P", family="paragraph")
    pstyle.addElement(style.TextProperties(attributes={"fontsize": "11pt"}))
    doc.styles.addElement(pstyle)

    ist = style.Style(name="I", family="text")
    ist.addElement(style.TextProperties(attributes={"fontstyle": "italic"}))
    doc.styles.addElement(ist)

    bst = style.Style(name="B", family="text")
    bst.addElement(style.TextProperties(attributes={"fontweight": "bold"}))
    doc.styles.addElement(bst)

    sst = style.Style(name="S", family="text")
    sst.addElement(style.TextProperties(attributes={"textlinethroughstyle": "solid"}))
    doc.styles.addElement(sst)

    H = struct["header"]

    # Sujet
    if H.get("sujet"):
        doc.text.addElement(text.P(stylename=h1, text=H["sujet"]))
    # Concours
    if H.get("concours"):
        doc.text.addElement(text.P(stylename=h2, text=H["concours"]))
    # Note
    if H.get("note"):
        doc.text.addElement(text.P(stylename=h3, text=f"Note : {H['note']}"))

    # Ligne vide
    doc.text.addElement(text.P(text=""))

    # Sections
    for label, paragraphs in struct["blocks"]:
        doc.text.addElement(text.P(stylename=hsec, text=label))
        for para in paragraphs:
            para = normalize_text(para)
            # On crée un paragraphe et on y ajoute des spans pour i/b/s
            p = text.P(stylename=pstyle)
            for i, line in enumerate(para.split("\n")):
                if i > 0:
                    p.addElement(text.LineBreak())
                for frag, style_key in parse_inline_md(line):
                    if style_key == "i":
                        span = text.Span(stylename=ist, text=frag)
                    elif style_key == "b":
                        span = text.Span(stylename=bst, text=frag)
                    elif style_key == "s":
                        span = text.Span(stylename=sst, text=frag)
                    else:
                        span = text.Span(text=frag)
                    p.addElement(span)
            doc.text.addElement(p)
        # Saut entre blocs
        doc.text.addElement(text.P(text=""))

    doc.save(str(odt_path))

# ---------- Point d’entrée ----------

def export_from_json(
    json_path: Union[str, Path],
    out_dir: Union[str, Path] = ".",
    basename: Optional[str] = None
):
    json_path = Path(json_path)
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    data = json.loads(json_path.read_text(encoding="utf-8"))
    struct = extract_structure(data)

    base = basename or (data.get("sujet") or json_path.stem)
    safe = re.sub(r"[^\w\-]+", "_", base).strip("_") or "dissertation"

    md_path = out_dir / f"{safe}.md"
    docx_path = out_dir / f"{safe}.docx"
    odt_path = out_dir / f"{safe}.odt"

    export_markdown(struct, md_path)
    export_docx(struct, docx_path)
    export_odt(struct, odt_path)

    return md_path, docx_path, odt_path


if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser(description="Exporter (MD/DOCX/ODT) une dissertation CAPES depuis un JSON.")
    ap.add_argument("json_path", help="Chemin du fichier JSON")
    ap.add_argument("-o", "--out-dir", default=".", help="Dossier de sortie")
    ap.add_argument("-b", "--basename", default=None, help="Nom de base des fichiers de sortie")
    args = ap.parse_args()

    md, docx, odt = export_from_json(args.json_path, args.out_dir, args.basename)
    print("Fichiers générés :")
    print(" -", md)
    print(" -", docx)
    print(" -", odt)
