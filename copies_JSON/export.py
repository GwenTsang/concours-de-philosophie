import json
import re
from pathlib import Path
from typing import Dict, List, Tuple, Optional

# --------- Utils Markdown & normalisation ---------

def normalize_text(s: Optional[str]) -> str:
    if not s:
        return ""
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = s.replace("\\-", "-").replace("\\<\\<", "«").replace("\\>\\>", "»").replace("\\\"", "\"")
    s = re.sub(r"[ \t]+\n", "\n", s)
    return s.strip("\n")

Token = Tuple[str, Optional[str]]
_inline_re = re.compile(r"(~~.+?~~|\*\*.+?\*\*|\*.+?\*)", flags=re.DOTALL)

def parse_inline_md(text: str) -> List[Token]:
    tokens: List[Token] = []
    pos = 0
    for m in _inline_re.finditer(text):
        if m.start() > pos:
            tokens.append((text[pos:m.start()], None))
        frag = m.group(0)
        if frag.startswith("**") and frag.endswith("**"):
            tokens.append((frag[2:-2], "b"))
        elif frag.startswith("*") and frag.endswith("*"):
            tokens.append((frag[1:-1], "i"))
        elif frag.startswith("~~") and frag.endswith("~~"):
            tokens.append((frag[2:-2], "s"))
        else:
            tokens.append((frag, None))
        pos = m.end()
    if pos < len(text):
        tokens.append((text[pos:], None))
    return tokens

def md_to_plain_paragraphs(s: str) -> List[str]:
    s = normalize_text(s)
    if not s:
        return []
    blocks = re.split(r"\n\s*\n", s)
    out: List[str] = []
    for b in blocks:
        if not b:
            continue
        b = re.sub(r"^\n+|\n+$", "", b)
        lines = [ln.rstrip() for ln in b.split("\n")]
        block = "\n".join(lines)
        if block.strip():
            out.append(block)
    return out

# --------- Extraction du schéma ---------

def extract_structure(data: Dict) -> Dict:
    TAB = "\t"
    header = {
        "concours": " ".join([v for v in [data.get("niveau"), data.get("annee")] if v]),
        "note": data.get("note"),
        "sujet": data.get("sujet") or data.get("title") or data.get("titre"),
    }

    intro = md_to_plain_paragraphs(data.get("introduction") or "")
    annonce = normalize_text(data.get("annonce_du_plan") or "").strip("\n")
    if annonce:
        if intro:
            intro[-1] = intro[-1].rstrip() + f"\n{TAB}{annonce}"
        else:
            intro = [f"{TAB}{annonce}"]

    blocks_order = [
        ("Introduction", intro),
        ("I", data.get("partie_1")),
        ("Transition 1", data.get("transition_1")),
        ("II", data.get("partie_2")),
        ("Transition 2", data.get("transition_2")),
        ("III", data.get("partie_3")),
        ("Conclusion", data.get("conclusion")),
    ]

    blocks: List[Tuple[str, List[str]]] = []
    for label, raw in blocks_order:
        if isinstance(raw, list):
            pars = [p for p in raw if isinstance(p, str) and p.strip()]
        elif isinstance(raw, str):
            pars = md_to_plain_paragraphs(raw)
        else:
            pars = []
        if pars:
            blocks.append((label, pars))
    return {"header": header, "blocks": blocks}

# --------- Exports ---------

def export_markdown(struct: Dict, md_path: Path, rawtext: bool = False) -> None:
    H = struct["header"]
    lines: List[str] = [
        '<div align="center">',
        f"<h1>{H['sujet']}</h1>" if H.get("sujet") else "",
        f"<h2>{H['concours']}</h2>" if H.get("concours") else "",
        f"<h3>Note : {H['note']}</h3>" if H.get("note") else "",
        "</div>",
        ""
    ]
    for label, paragraphs in struct["blocks"]:
        is_transition = bool(re.match(r"Transition\s*\d+", label, flags=re.IGNORECASE))
        if not is_transition and not rawtext:
            lines.append(f'<div align="center"><h2>{label}</h2></div>\n')
        for p in paragraphs:
            lines.append(p)
            lines.append("")
        lines.append("")
    md_path.write_text("\n".join([l for l in lines if l is not None]).strip() + "\n", encoding="utf-8")

def _docx_add_line_with_leading_tabs(paragraph, line: str, style_key: Optional[str]):
    m = re.match(r"^\t+", line)
    tabs = len(m.group(0)) if m else 0
    content = line[tabs:]
    for _ in range(tabs):
        paragraph.add_run("\t")
    run = paragraph.add_run(content)
    if style_key == "i": run.italic = True
    elif style_key == "b": run.bold = True
    elif style_key == "s": run.font.strike = True

def export_docx(struct: Dict, docx_path: Path, rawtext: bool = False) -> None:
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ModuleNotFoundError:
        raise ImportError("Installez python-docx : pip install python-docx")

    doc = Document()
    H = struct["header"]

    if H.get("sujet"):
        p = doc.add_paragraph(); r = p.add_run(H["sujet"]); r.bold = True; r.font.size = Pt(22)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if H.get("concours"):
        p = doc.add_paragraph(); r = p.add_run(H["concours"]); r.bold = True; r.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if H.get("note"):
        p = doc.add_paragraph(); r = p.add_run(f"Note : {H['note']}"); r.bold = True; r.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    for label, paragraphs in struct["blocks"]:
        is_transition = bool(re.match(r"Transition\s*\d+", label, flags=re.IGNORECASE))
        if not is_transition and not rawtext:
            h = doc.add_paragraph(); hr = h.add_run(label); hr.bold = True; hr.font.size = Pt(14)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for para in paragraphs:
            para = normalize_text(para)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(0.75)
            for i, line in enumerate(para.split("\n")):
                if i > 0: p.add_run("\n")
                tokens = parse_inline_md(line)
                if not tokens:
                    _docx_add_line_with_leading_tabs(p, line, None)
                else:
                    for j, (frag, style) in enumerate(tokens):
                        if j == 0: _docx_add_line_with_leading_tabs(p, frag, style)
                        else:
                            r = p.add_run(frag)
                            if style == "i": r.italic = True
                            elif style == "b": r.bold = True
                            elif style == "s": r.font.strike = True
        doc.add_paragraph("")
    doc.save(str(docx_path))

def export_odt(struct: Dict, odt_path: Path, rawtext: bool = False) -> None:
    from odf.opendocument import OpenDocumentText
    from odf import text, style
    doc = OpenDocumentText()

    def pstyle(name, props=None, tprops=None):
        st = style.Style(name=name, family="paragraph")
        if props: st.addElement(style.ParagraphProperties(attributes=props))
        if tprops: st.addElement(style.TextProperties(attributes=tprops))
        doc.styles.addElement(st); return st

    h1 = pstyle("H1", {"textalign": "center"}, {"fontsize": "22pt", "fontweight": "bold"})
    h2 = pstyle("H2", {"textalign": "center"}, {"fontsize": "16pt", "fontweight": "bold"})
    h3 = pstyle("H3", {"textalign": "center"}, {"fontsize": "14pt", "fontweight": "bold"})
    hsec = pstyle("HSEC", {"textalign": "center"}, {"fontsize": "14pt", "fontweight": "bold"})
    body = pstyle("P", {"textindent": "0.75cm", "textalign": "justify"}, {"fontsize": "11pt"})

    ist = style.Style(name="I", family="text"); ist.addElement(style.TextProperties(attributes={"fontstyle":"italic"})); doc.styles.addElement(ist)
    bst = style.Style(name="B", family="text"); bst.addElement(style.TextProperties(attributes={"fontweight":"bold"})); doc.styles.addElement(bst)
    sst = style.Style(name="S", family="text"); sst.addElement(style.TextProperties(attributes={"textlinethroughstyle":"solid"})); doc.styles.addElement(sst)

    H = struct["header"]
    if H.get("sujet"): doc.text.addElement(text.P(stylename=h1, text=H["sujet"]))
    if H.get("concours"): doc.text.addElement(text.P(stylename=h2, text=H["concours"]))
    if H.get("note"): doc.text.addElement(text.P(stylename=h3, text=f"Note : {H['note']}"))
    doc.text.addElement(text.P(text=""))

    for label, paragraphs in struct["blocks"]:
        is_transition = bool(re.match(r"Transition\s*\d+", label, flags=re.IGNORECASE))
        doc.text.addElement(text.P(stylename=hsec, text=label) if (not is_transition and not rawtext) else text.P(text=""))
        for para in paragraphs:
            for line in normalize_text(para).split("\n"):
                p = text.P(stylename=body)
                m = re.match(r"^\t+", line)
                tabs = len(m.group(0)) if m else 0
                rest = line[tabs:]
                for _ in range(tabs): p.addElement(text.Tab())
                for frag, sk in parse_inline_md(rest):
                    if sk == "i": span = text.Span(stylename=ist, text=frag)
                    elif sk == "b": span = text.Span(stylename=bst, text=frag)
                    elif sk == "s": span = text.Span(stylename=sst, text=frag)
                    else: span = text.Span(text=frag)
                    p.addElement(span)
                doc.text.addElement(p)
        doc.text.addElement(text.P(text=""))
    doc.save(str(odt_path))

# --------- Batch dossier uniquement ---------

def export_folder(folder: Path, out_dir: Path, rawtext: bool = False, recursive: bool = False):
    if not folder.is_dir():
        raise NotADirectoryError(f"« {folder} » n'est pas un dossier existant.")
    out_dir.mkdir(parents=True, exist_ok=True)

    iterator = folder.rglob("*.json") if recursive else folder.glob("*.json")
    json_files = sorted(p for p in iterator if p.is_file())

    if not json_files:
        print(f"Aucun JSON trouvé dans {folder}" + (" (récursif)" if recursive else ""))
        return

    for jp in json_files:
        try:
            data = json.loads(jp.read_text(encoding="utf-8"))
            struct = extract_structure(data)
            base = (data.get("sujet") or jp.stem)
            safe = re.sub(r"[^\w\-]+", "_", base).strip("_") or "dissertation"
            export_markdown(struct, out_dir / f"{safe}.md", rawtext=rawtext)
            export_docx(struct, out_dir / f"{safe}.docx", rawtext=rawtext)
            export_odt(struct, out_dir / f"{safe}.odt", rawtext=rawtext)
            print(f"OK: {jp}")
        except Exception as e:
            print(f"ERREUR: {jp} -> {e}")

# --------- CLI ---------

if __name__ == "__main__":
    import argparse
    script_dir = Path(__file__).resolve().parent  # <-- dossier du script

    ap = argparse.ArgumentParser(
        description="Exporter (MD/DOCX/ODT) des dissertations CAPES depuis un dossier de JSON."
    )
    # 'folder' devient optionnel; par défaut: dossier du script
    ap.add_argument(
        "folder",
        nargs="?",
        default=str(script_dir),
        help=f"Dossier contenant des fichiers JSON (par défaut : {script_dir})"
    )
    ap.add_argument("-o", "--out-dir", default=".", help="Dossier de sortie")
    ap.add_argument("--recursive", action="store_true", help="Inclure les sous-dossiers")
    ap.add_argument("-rawtext", "--rawtext", action="store_true",
                    help="Sans balises de sections (Intro, I, II, III, Conclusion)")
    args = ap.parse_args()

    export_folder(
        Path(args.folder),
        Path(args.out_dir),
        rawtext=args.rawtext,
        recursive=args.recursive
    )
